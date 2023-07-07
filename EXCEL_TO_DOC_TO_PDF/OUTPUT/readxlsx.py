import pandas as pd
from docx2pdf import convert
import openpyxl
from docx import Document
from  docx.shared import Pt

def RUN():
    doc = Document('DOCC.docx')
    df = pd.read_excel('การลงทะเบียนเข้าร่วมกิจกรรม (การตอบกลับ).xlsx')  # , usecols='')
    tt = ["", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "","", ""]
    # print(df.iloc[1][1])
    new_data = []
    number = 23
    for x in range(number-1):
        for i in range(29):
            f = str(df.iloc[x][i])
            tt[i] = f

        for i in range(29):
            tt[i] = tt[i].split("\n")
        for i in range(29):
            tt[i] = str(tt[i][0])

        LEN = len(tt[2])
        if tt[2][LEN-1] == "1":
            doc.tables[0].cell(0, 1).paragraphs[0].text = tt[2][0:LEN-1]
        else:
            doc.tables[0].cell(0, 1).paragraphs[0].text = tt[2]

        doc.tables[0].cell(1, 1).paragraphs[0].text = "แผนกวิชา" + tt[3]
        doc.tables[0].cell(2, 2).paragraphs[1].text = tt[4]
        doc.tables[0].cell(2, 2).paragraphs[2].text = tt[6]
        doc.tables[0].cell(2, 3).paragraphs[1].text = "รหัสนักศึกษา " + tt[5]
        doc.tables[0].cell(2, 3).paragraphs[2].text = "รหัสนักศึกษา " + tt[7]
        if tt[8] != "nan" and tt[9] != "nan":
            pass
            doc.tables[0].cell(2, 1).paragraphs[3].text = "2. ชื่อ-สกุล"
            doc.tables[0].cell(2, 2).paragraphs[3].text = tt[8]
            doc.tables[0].cell(2, 3).paragraphs[3].text = "รหัสนักศึกษา " +tt[9]

        elif tt[8] == "nan" and tt[9] == "nan":
            doc.tables[0].cell(2, 1).paragraphs[3].text =""
            doc.tables[0].cell(2, 2).paragraphs[3].text = ""
            doc.tables[0].cell(2, 3).paragraphs[3].text = ""




        doc.tables[0].cell(3, 0).paragraphs[1].text = "- " + tt[10]
        doc.tables[0].cell(4, 0).paragraphs[1].text = "ภาคเรียนที่ " + tt[11] + "   ปีการศึกษา " + tt[12]
        doc.tables[0].cell(5, 0).paragraphs[1].text = tt[13]
        doc.tables[0].cell(6, 0).paragraphs[1].text = "1. เพื่อออกแบบเครื่อง" + tt[14]
        doc.tables[0].cell(6, 0).paragraphs[2].text = "2. เพื่อสร้างเครื่อง" + tt[15]
        doc.tables[0].cell(6, 0).paragraphs[3].text = "3. เพื่อทดสอบเครื่อง" + tt[16]
        doc.tables[0].cell(6, 0).paragraphs[4].text = "4. เพื่อหาประสิทธิภาพเครื่อง" + tt[17]
        doc.tables[0].cell(7, 0).paragraphs[1].text = tt[18]
        doc.tables[0].cell(9, 0).paragraphs[1].text = tt[25] + " บาท"
        doc.tables[0].cell(10, 0).paragraphs[1].text = "ได้เครื่อง" + tt[26]
        doc.tables[1].cell(1, 0).paragraphs[0].text = "ผลโครงงานได้เครื่อง" + tt[27]

        new_data.append(doc.tables[0].cell(0, 1).paragraphs[0])
        new_data.append(doc.tables[0].cell(1, 1).paragraphs[0])
        new_data.append(doc.tables[0].cell(3, 0).paragraphs[1])
        new_data.append(doc.tables[0].cell(4, 0).paragraphs[1])
        new_data.append(doc.tables[0].cell(5, 0).paragraphs[1])
        new_data.append(doc.tables[0].cell(6, 0).paragraphs[1])
        new_data.append(doc.tables[0].cell(6, 0).paragraphs[2])
        new_data.append(doc.tables[0].cell(6, 0).paragraphs[3])
        new_data.append(doc.tables[0].cell(6, 0).paragraphs[4])
        new_data.append(doc.tables[0].cell(7, 0).paragraphs[1])
        new_data.append(doc.tables[0].cell(9, 0).paragraphs[1])
        new_data.append(doc.tables[0].cell(10, 0).paragraphs[1])
        new_data.append(doc.tables[0].cell(0, 1).paragraphs[0])
        new_data.append(doc.tables[0].cell(1, 1).paragraphs[0])
        new_data.append(doc.tables[0].cell(2, 2).paragraphs[1])
        new_data.append(doc.tables[0].cell(2, 3).paragraphs[1])
        new_data.append(doc.tables[1].cell(1, 0).paragraphs[0])
        new_data.append(doc.tables[0].cell(2, 2).paragraphs[2])
        new_data.append(doc.tables[0].cell(2, 3).paragraphs[2])

        new_data.append(doc.tables[0].cell(2, 1).paragraphs[3])
        new_data.append(doc.tables[0].cell(2, 2).paragraphs[3])
        new_data.append(doc.tables[0].cell(2, 3).paragraphs[3])


        for data in new_data:
            data.runs[0].font.name = 'TH SarabunPSK'
            data.runs[0].font.size = Pt(16)




        doc.save("C:\\Users\\ASUS\\Desktop\\EXCEL_TO_DOC\\OUTPUT\\DOCX\\"+tt[2]+'.docx')
        doc.save("C:\\Users\\ASUS\\Desktop\\EXCEL_TO_DOC\\OUTPUT\\PDF\\"+tt[2]+'.pdf')

        partout = "C:\\Users\\ASUS\\Desktop\\EXCEL_TO_DOC\\OUTPUT\\PDF\\"+tt[2]+'.pdf'
        partin = "C:\\Users\\ASUS\\Desktop\\EXCEL_TO_DOC\\OUTPUT\\DOCX\\"+tt[2]+'.docx'
        convert(partin,partout)

        #doc.save(tt[2] + '.pdf')
        print(tt[2])






RUN()




